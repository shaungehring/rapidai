"""Document loaders for RAG."""

from pathlib import Path
from typing import Union

from ..exceptions import DocumentLoaderError
from ..types import Document
from .base import BaseDocumentLoader


class PDFLoader(BaseDocumentLoader):
    """Load PDF documents."""

    async def load(self, source: Union[str, Path]) -> Document:
        """Load a PDF document.

        Args:
            source: File path to PDF

        Returns:
            Loaded document
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise DocumentLoaderError(
                "pypdf not installed. Install with: pip install rapidai-framework[rag]"
            )

        path = Path(source)
        if not path.exists():
            raise DocumentLoaderError(f"File not found: {path}")

        try:
            reader = PdfReader(path)

            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n\n"

            metadata = {
                "source": str(path),
                "type": "pdf",
                "pages": len(reader.pages),
                "filename": path.name,
            }

            return Document(content=content.strip(), metadata=metadata)
        except Exception as e:
            raise DocumentLoaderError(f"Failed to load PDF {path}: {str(e)}")


class DOCXLoader(BaseDocumentLoader):
    """Load DOCX documents."""

    async def load(self, source: Union[str, Path]) -> Document:
        """Load a DOCX document.

        Args:
            source: File path to DOCX

        Returns:
            Loaded document
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise DocumentLoaderError(
                "python-docx not installed. Install with: pip install rapidai-framework[rag]"
            )

        path = Path(source)
        if not path.exists():
            raise DocumentLoaderError(f"File not found: {path}")

        try:
            doc = DocxDocument(path)

            content = "\n\n".join([para.text for para in doc.paragraphs if para.text])

            metadata = {
                "source": str(path),
                "type": "docx",
                "filename": path.name,
            }

            return Document(content=content, metadata=metadata)
        except Exception as e:
            raise DocumentLoaderError(f"Failed to load DOCX {path}: {str(e)}")


class TextLoader(BaseDocumentLoader):
    """Load plain text files."""

    async def load(self, source: Union[str, Path]) -> Document:
        """Load a text file.

        Args:
            source: File path to text file

        Returns:
            Loaded document
        """
        path = Path(source)
        if not path.exists():
            raise DocumentLoaderError(f"File not found: {path}")

        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()

            metadata = {
                "source": str(path),
                "type": "txt",
                "filename": path.name,
            }

            return Document(content=content, metadata=metadata)
        except Exception as e:
            raise DocumentLoaderError(f"Failed to load text file {path}: {str(e)}")


class MarkdownLoader(TextLoader):
    """Load Markdown files."""

    async def load(self, source: Union[str, Path]) -> Document:
        """Load a Markdown file.

        Args:
            source: File path to Markdown file

        Returns:
            Loaded document
        """
        doc = await super().load(source)
        doc.metadata["type"] = "markdown"
        return doc


class HTMLLoader(BaseDocumentLoader):
    """Load HTML documents."""

    async def load(self, source: Union[str, Path]) -> Document:
        """Load an HTML file.

        Args:
            source: File path to HTML file

        Returns:
            Loaded document
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise DocumentLoaderError(
                "beautifulsoup4 not installed. Install with: pip install rapidai-framework[rag]"
            )

        path = Path(source)
        if not path.exists():
            raise DocumentLoaderError(f"File not found: {path}")

        try:
            with open(path, "r", encoding="utf-8") as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            content = soup.get_text(separator="\n\n")

            # Clean up whitespace
            lines = (line.strip() for line in content.splitlines())
            content = "\n\n".join(line for line in lines if line)

            metadata = {
                "source": str(path),
                "type": "html",
                "filename": path.name,
                "title": soup.title.string if soup.title else None,
            }

            return Document(content=content.strip(), metadata=metadata)
        except Exception as e:
            raise DocumentLoaderError(f"Failed to load HTML {path}: {str(e)}")


def DocumentLoader(source: Union[str, Path]) -> BaseDocumentLoader:
    """Factory function to create document loader based on file extension.

    Args:
        source: File path

    Returns:
        Appropriate document loader instance

    Example:
        ```python
        loader = DocumentLoader("document.pdf")
        doc = await loader.load("document.pdf")
        ```
    """
    path = Path(source)
    ext = path.suffix.lower()

    loaders = {
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
        ".txt": TextLoader,
        ".md": MarkdownLoader,
        ".markdown": MarkdownLoader,
        ".html": HTMLLoader,
        ".htm": HTMLLoader,
    }

    loader_class = loaders.get(ext)
    if not loader_class:
        raise DocumentLoaderError(
            f"Unsupported file type: {ext}. "
            f"Supported types: {', '.join(loaders.keys())}"
        )

    return loader_class()
